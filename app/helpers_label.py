from datetime import datetime

from flask import (
    current_app,
)

from docx import Document
from docx.shared import Pt, Mm, Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

# def _add_label_border(paragraph, border_color="000000", border_size="12"):
#     p = paragraph._element
#     pPr = p.get_or_add_pPr()
#     pbdr = OxmlElement('w:pBdr')
#     for side in ['top', 'left', 'bottom', 'right']:
#         border = OxmlElement(f'w:{side}')
#         border.set(qn('w:val'), 'single')
#         border.set(qn('w:sz'), border_size)  # Thickness
#         border.set(qn('w:space'), '4')       # Space between text and border
#         border.set(qn('w:color'), border_color)
#         pbdr.append(border)
#     pPr.append(pbdr)


class SpecimenLabel(object):
    _doc = None
    fonts = {
        'l': Pt(12),
        'm': Pt(10),
        's': Pt(8),
    }

    def __init__(self, type_='docx'):
        if type_ == 'docx':
            self._doc = Document()

    def set_docx_layout(self):
        # make A4, 2-column layout
        section = self._doc.sections[0]
        sectPr = section._sectPr
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')  # Space between columns (0.5 inch in twips)
        sectPr.append(cols)

        # Set up margins
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # for ref
        # section = doc.sections[0]
        # section.page_width = Mm(210)
        # section.page_height = Mm(297)
        # section.top_margin = Mm(15)
        # section.bottom_margin = Mm(15)
        # section.left_margin = Mm(15)
        # section.right_margin = Mm(15)
        # section.column_width = Mm(85)
        # #section.start_type

        # # set_2_columns
        # section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')

        style = self._doc.styles['Normal']
        style.paragraph_format.line_spacing = 1

    def add_docx_text(self, content='', size='m', styles=[]):
        p = self._doc.add_paragraph()
        if 'center' in styles:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if content:
            run = p.add_run(content)
            run.font.size = self.fonts[size]

            if 'bold' in styles:
                run.font.bold = True

        return p

    def add_docx_text_between(self, left, right, size='m'):
        p = self.add_docx_text(left, size)

        tab_runs = p.add_run('\t')
        p.add_run(right).font.size = self.fonts[size]
        # Add right tab stop for alignment
        p.paragraph_format.tab_stops.add_tab_stop(Inches(3.5), WD_ALIGN_PARAGRAPH.RIGHT)

    def as_docx(self):
        return self._doc

    def get_label_text(self, entity):
        record = entity['record']
        unit = entity['unit']

        header = ''
        if country := record.get_named_area('COUNTRY'):
            if country.id == 1311:
                header = 'BOTANICAL INVENTORY OF TAIWAN'
            else:
                header = f'PLANTS OF {country.name_en.upper()}'

        ids = []
        for i in record.identifications:
            taxon_family = ''
            taxon_species = ''
            if i.taxon:
                if family := i.taxon.get_higher_taxon('family'):
                    taxon_family = family.full_scientific_name
                taxon_species = i.taxon.display_name
                id_data = {
                    'taxon': [taxon_family, taxon_species],
                }
            if x := i.identifier:
                id_data.update({'identifier': x.full_name_en or x.full_name})
            if x := i.date:
                id_data.update({'date': i.get_date_display('%b. %d, %Y')})
            ids.append(id_data)

        locality_list = []
        na = []
        for k, v in record.get_named_area_map().items():
            if v and k != 'COUNTRY':
                na.append(v.named_area.display_name)
        if len(na) > 0:
            locality_list.append(' ,'.join(na))
        if x := record.locality_text:
            locality_list.append(x)
        if x := record.locality_text_en:
            locality_list.append(x)

        lonlat = ''
        elev = ''
        if coordinates := record.get_coordinates('dms'):
            lonlat = f"{coordinates['x_label']}, {coordinates['y_label']}"
        if record.altitude:
            elev = f"Elev. ca. {record.altitude}"
            if x := record.altitude2:
                elev = f'{elev}-{x}'
            elev = f'{elev} m'

        data = {
            'identifications': ids,
            'layout': {
                'header': header,
                'footer': ['中央研究院植物標本館', 'Herbarium, Academia Sinica, Taipei (HAST)'],
            },
            'event': {
                'collector': '',
                'companion': '',
                'date': '',
                'location': {
                    'lonlat': lonlat,
                    'elev': elev,
                    'locality': locality_list,
                }
            },
            'unit': {},
        }

        collector = ''
        if x := record.collector_id:
            collector = record.collector.full_name_en or record.collector.full_name
        if y := record.field_number:
            collector = f'{collector} {y}'
            data['event']['collector'] = collector
        data['event']['people'] = collector
        if x := record.collect_date:
            data['event']['date'] = x.strftime('%b. %d, %Y')
        if x := record.companion_list:
            data['event']['companion'] = x

        # only HAST rules now
        notes = []
        appendix = []
        type_status = ''
        if entity['type'] == 'unit':
            unit_data = {}
            unit = entity['unit']
            if x := unit.notes:
                notes.append(x)
            if x := unit.get_annotation('add-char'):
                notes.append(unit.get_annotation('add-char').value)

            if x := unit.get_annotation('greenhouse'):
                if x.value:
                    appendix.append('Plants of this collection were brought back for cultivation.')

            if p_date := unit.get_annotation('greenhouse_pressed_date'):
                d = ''
                try:
                    d = datetime.strptime(p_date.value, '%Y-%m-%d').strftime('%b. %d, %Y')
                    appendix.append(f'This specimen was pressed on {d}')
                except:
                    current_app.logger.error('greenhouse_pressed_date date format error')

            if unit.type_status:
                data['unit']['type_status'] = f'{unit.type_status.upper()} of {unit.typified_name}'

        for i in entity['assertion_display_list']:
            notes.append(i)

        if x := record.field_note:
            notes.append(x)
        if x := record.field_note_en:
            notes.append(x)

        if len(notes):
            data['notes'] = notes

        if len(appendix):
            data['unit']['appendix'] = appendix

        return data

def make_print_docx(items):
    labels = SpecimenLabel('docx')
    labels.set_docx_layout()

    for item in items:
        label_data = labels.get_label_text(item)
        #_create_label(doc, label_info)
        #doc.add_paragraph()  # Spacing between labels

        if len(label_data['identifications']) > 1:
            for k, v in enumerate(label_data['identifications'][1:]):
                t = v['taxon'][0]
                if x := v['taxon'][1]:
                    t = f'{t}\n{x}'
                labels.add_docx_text(t, 'm', ['bold'])

                left = f"Det. by {v['identifier']}" if v.get('identifier') else ''
                right = v['date'] if v.get('date') else ''
                labels.add_docx_text_between(left, right)

        labels.add_docx_text(label_data['layout']['header'], 'l', ['bold', 'center'])
        if len(label_data['identifications']) > 0:
            if first_id := label_data['identifications'][0]:
                t = first_id['taxon'][0].upper()
                if species := first_id['taxon'][1]:
                    t = f'{t}\n{species}'
                labels.add_docx_text(t, 'm', ['bold'])


        locality = '\n'.join(label_data['event']['location']['locality'])
        labels.add_docx_text(locality)

        labels.add_docx_text_between(label_data['event']['location']['lonlat'], label_data['event']['location']['elev'])

        if notes := label_data.get('notes'):
            notes = '\n'.join(notes)
            labels.add_docx_text(notes)

        if label_data['unit']:
            if x := label_data['unit'].get('type_status'):
                labels.add_docx_text(x)

        if label_data['event']['collector'] or label_data['event']['date']:
            labels.add_docx_text_between(label_data['event']['collector'], label_data['event']['date'])
        if x:= label_data['event'].get('companion'):
            s = '& ' + ', '.join(x)
            labels.add_docx_text(s)

        footer = '\n'.join(label_data['layout']['footer'])
        labels.add_docx_text(footer, 's', ['bold', 'center'])

        if x := label_data['unit'].get('appendix'):
            appendix = '\n'.join(x)
            labels.add_docx_text(appendix, 's', ['center'])

        # Add space between labels
        labels.add_docx_text('')

    return labels.as_docx()
